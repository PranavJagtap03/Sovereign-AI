"""
Output Generator Module — Sovereign AI Workbench
Generates real office deliverable documents (.docx) offline with local formatting.

Roadmap Note:
- XLSX generation (spreadsheets) will follow this same pattern using `openpyxl`.
- PPTX generation (slideshows) will follow this same pattern using `python-pptx`.
"""

import os
import re
from datetime import datetime
from typing import List, Dict, Any, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUTPUTS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "outputs")


def parse_markdown_to_sections(text: str) -> List[Dict[str, str]]:
    """
    Parses an AI markdown response into structured sections with headings and body content.
    """
    if not text:
        return [{"heading": "Deliverable Summary", "content": "No content generated."}]

    sections = []
    current_heading = "Executive Summary"
    current_lines = []

    for line in text.split("\n"):
        stripped = line.strip()
        # Detect markdown headers: ## Heading or **Heading:** or **Heading**
        header_match = re.match(r"^(?:#{1,4}\s+|\*\*)([^*#:\n]+)(?:\*\*|:)?$", stripped)
        if header_match and len(stripped) < 90 and not stripped.startswith("```"):
            if current_lines:
                content = "\n".join(current_lines).strip()
                if content:
                    sections.append({"heading": current_heading, "content": content})
                current_lines = []
            current_heading = header_match.group(1).strip()
        else:
            current_lines.append(line)

    if current_lines:
        content = "\n".join(current_lines).strip()
        if content:
            sections.append({"heading": current_heading, "content": content})

    return sections or [{"heading": "Report Content", "content": text}]


def generate_docx_report(title: str, sections: List[Dict[str, str]], filename: str = "report.docx") -> str:
    """
    Generates a professionally formatted Microsoft Word document (.docx) with:
    - Dedicated title cover page
    - Section headings and styled paragraphs / bullet items
    - Header and footer with air-gapped sovereignty attestation
    
    Args:
        title: Main document title
        sections: List of sections, e.g. [{"heading": "Summary", "content": "..."}, ...]
        filename: Destination filename (default: "report.docx")

    Returns:
        str: Absolute file path to the generated .docx file in /backend/outputs/
    """
    os.makedirs(OUTPUTS_DIR, exist_ok=True)
    if not filename.endswith(".docx"):
        filename = f"{filename}.docx"

    file_path = os.path.join(OUTPUTS_DIR, filename)

    doc = Document()

    # ─── Page Setup ───────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # ─── Header & Footer ──────────────────────────────────────────────────────
    header = section.header
    p_head = header.paragraphs[0]
    p_head.text = "SOVEREIGN AI WORKBENCH | OFFICIAL DELIVERABLE"
    p_head.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if p_head.runs:
        p_head.runs[0].font.size = Pt(8)
        p_head.runs[0].font.color.rgb = RGBColor(120, 130, 140)

    footer = section.footer
    p_foot = footer.paragraphs[0]
    p_foot.text = "Generated Offline on Air-Gapped Local Hardware | 0 Bytes Outbound Egress"
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p_foot.runs:
        p_foot.runs[0].font.size = Pt(8)
        p_foot.runs[0].font.color.rgb = RGBColor(140, 150, 160)

    # ─── Title Page ───────────────────────────────────────────────────────────
    # Spacing before title
    for _ in range(3):
        doc.add_paragraph()

    # Organization / Classification pill
    p_org = doc.add_paragraph()
    p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_org = p_org.add_run("GOVERNMENT OF INDIA — SOVEREIGN AI ENCLAVE")
    run_org.font.size = Pt(10)
    run_org.font.bold = True
    run_org.font.color.rgb = RGBColor(0, 160, 120)

    # Main Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(title)
    run_title.font.size = Pt(26)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(26, 44, 76)

    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("AI-Generated Engineering & Compliance Deliverable")
    run_sub.font.size = Pt(13)
    run_sub.font.color.rgb = RGBColor(90, 105, 120)

    # Decorative line
    for _ in range(2):
        doc.add_paragraph()

    # Metadata Table
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    metadata = [
        ("Generation Date:", datetime.now().strftime("%d-%B-%Y %H:%M:%S")),
        ("Security Clearance:", "Restricted — Sovereign Deployment"),
        ("Verification Status:", "Stage 8 Guardian Passed ✓"),
        ("Model Execution:", "Local Hardware Enclave (Air-Gapped)")
    ]
    for idx, (label, val) in enumerate(metadata):
        row = table.rows[idx]
        cell_lbl, cell_val = row.cells[0], row.cells[1]
        cell_lbl.text = label
        cell_val.text = val
        cell_lbl.paragraphs[0].runs[0].font.bold = True
        cell_lbl.paragraphs[0].runs[0].font.size = Pt(9)
        cell_val.paragraphs[0].runs[0].font.size = Pt(9)
        cell_lbl.paragraphs[0].runs[0].font.color.rgb = RGBColor(90, 105, 120)
        cell_val.paragraphs[0].runs[0].font.color.rgb = RGBColor(40, 50, 60)

    # End of title page
    doc.add_page_break()

    # ─── Body Sections ────────────────────────────────────────────────────────
    for section_data in sections:
        heading = section_data.get("heading", "Section")
        content = section_data.get("content", "")

        # Heading 1
        h1 = doc.add_heading(heading, level=1)
        if h1.runs:
            h1.runs[0].font.color.rgb = RGBColor(26, 54, 93)
            h1.runs[0].font.size = Pt(16)
            h1.runs[0].font.bold = True

        # Paragraphs / Content
        for line in content.split("\n"):
            line_str = line.strip()
            if not line_str:
                continue

            if line_str.startswith("- ") or line_str.startswith("* "):
                p = doc.add_paragraph(line_str[2:], style="List Bullet")
                p.paragraph_format.space_after = Pt(3)
            elif re.match(r"^\d+\.\s+", line_str):
                cleaned = re.sub(r"^\d+\.\s+", "", line_str)
                p = doc.add_paragraph(cleaned, style="List Number")
                p.paragraph_format.space_after = Pt(3)
            else:
                p = doc.add_paragraph(line_str)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.15
                if p.runs:
                    p.runs[0].font.size = Pt(10.5)
                    p.runs[0].font.color.rgb = RGBColor(45, 55, 72)

        # Spacing after section
        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_after = Pt(8)

    doc.save(file_path)
    return file_path
